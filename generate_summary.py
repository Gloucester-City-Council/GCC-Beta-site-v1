from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Colour palette ────────────────────────────────────────────────────────────
GCC_BLUE   = RGBColor(0x00, 0x54, 0xA4)   # --primary-base
DARK_BLUE  = RGBColor(0x00, 0x26, 0x50)   # --primary-darker
TEAL       = RGBColor(0x0A, 0x9E, 0x9E)   # --secondary-base
MID_GREY   = RGBColor(0x6C, 0x75, 0x7D)
LIGHT_GREY = RGBColor(0xF1, 0xF3, 0xF5)
BLACK      = RGBColor(0x1A, 0x1A, 0x2E)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helper: shade a table cell ────────────────────────────────────────────────
def shade_cell(cell, hex_colour):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_colour)
    tcPr.append(shd)

# ── Helper: add a coloured horizontal rule ────────────────────────────────────
def add_rule(doc, colour_hex='0054A4', thickness=12):
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    btm  = OxmlElement('w:bottom')
    btm.set(qn('w:val'),   'single')
    btm.set(qn('w:sz'),    str(thickness))
    btm.set(qn('w:space'), '1')
    btm.set(qn('w:color'), colour_hex)
    pBdr.append(btm)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    return p

# ── Helper: styled heading ────────────────────────────────────────────────────
def add_heading(doc, text, level=1, colour=None, space_before=18, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = colour or GCC_BLUE
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = colour or DARK_BLUE
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = colour or BLACK
    return p

# ── Helper: body paragraph ────────────────────────────────────────────────────
def add_body(doc, text, space_after=4, italic=False, colour=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.italic = italic
    if colour:
        run.font.color.rgb = colour
    return p

# ── Helper: tick bullet ───────────────────────────────────────────────────────
def add_tick(doc, text, done=True):
    p   = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_after  = Pt(2)
    mark = '✓  ' if done else '○  '
    run  = p.add_run(mark)
    run.font.color.rgb = TEAL if done else MID_GREY
    run.font.size = Pt(10.5)
    run.bold = done
    run2 = p.add_run(text)
    run2.font.size = Pt(10.5)
    return p

# ── Helper: todo bullet ───────────────────────────────────────────────────────
def add_todo(doc, label, priority, description):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'{label}  ')
    run.bold = True
    run.font.size = Pt(10.5)
    col = GCC_BLUE if priority == 'High' else (TEAL if priority == 'Medium' else MID_GREY)
    run.font.color.rgb = col
    run2 = p.add_run(description)
    run2.font.size = Pt(10.5)

# ═════════════════════════════════════════════════════════════════════════════
# COVER BLOCK
# ═════════════════════════════════════════════════════════════════════════════

# Blue header band via 1-row table
tbl = doc.add_table(rows=1, cols=1)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = tbl.cell(0, 0)
shade_cell(cell, '0054A4')
cell.width = Inches(6.5)
p = cell.paragraphs[0]
p.paragraph_format.space_before = Pt(18)
p.paragraph_format.space_after  = Pt(18)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run('Gloucester City Council')
r.font.color.rgb = WHITE
r.font.size = Pt(11)
r.bold = True
p2 = cell.add_paragraph()
p2.paragraph_format.space_before = Pt(2)
p2.paragraph_format.space_after  = Pt(18)
p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
r2 = p2.add_run('HMO Beta Website — Development Summary')
r2.font.color.rgb = WHITE
r2.font.size = Pt(20)
r2.bold = True
p3 = cell.add_paragraph()
p3.paragraph_format.space_before = Pt(0)
p3.paragraph_format.space_after  = Pt(18)
p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
r3 = p3.add_run(f'Session date: {datetime.date.today().strftime("%-d %B %Y")}   ·   Branch: claude/elegant-knuth-v8YaS   ·   PR: #43')
r3.font.color.rgb = RGBColor(0xCC, 0xE0, 0xFF)
r3.font.size = Pt(9.5)

doc.add_paragraph()

# ── Intro ─────────────────────────────────────────────────────────────────────
add_body(doc,
    'This document summarises the review and development work completed on the HMO '
    'section of the Gloucester City Council beta website during this session. It covers '
    'every issue that was investigated and resolved, the decisions made, and the work '
    'that remains outstanding. It is intended as a handover reference for the wider team.',
    space_after=8)

add_rule(doc)

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 1 — WHAT WAS REVIEWED
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '1.  What was reviewed', level=1)
add_body(doc,
    'A full audit of the HMO service pages was carried out against a list of issues '
    'covering design, navigation, functionality, UI components and accessibility. '
    'The pages reviewed were:')

pages = [
    ('index.html',      'HMO service landing page'),
    ('landlords.html',  'HMO landlord and manager information'),
    ('licensing.html',  'HMO licensing in Gloucester'),
    ('residents.html',  'Living in an HMO — your rights and what to expect'),
    ('standards.html',  'HMO standards and safety requirements'),
    ('planning.html',   'Planning permission for Houses in Multiple Occupation (HMO)'),
    ('apply.html',      'Apply for an HMO licence'),
    ('faq.html',        'HMO frequently asked questions (now a navigation page)'),
]
tbl = doc.add_table(rows=len(pages)+1, cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
# Header row
for i, txt in enumerate(['File', 'Page title']):
    cell = tbl.cell(0, i)
    shade_cell(cell, '002650')
    run = cell.paragraphs[0].add_run(txt)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)
for r, (f, t) in enumerate(pages, 1):
    tbl.cell(r, 0).paragraphs[0].add_run(f).font.size = Pt(10)
    tbl.cell(r, 1).paragraphs[0].add_run(t).font.size = Pt(10)
    if r % 2 == 0:
        shade_cell(tbl.cell(r, 0), 'F1F3F5')
        shade_cell(tbl.cell(r, 1), 'F1F3F5')

doc.add_paragraph()
add_body(doc,
    'The shared stylesheet (styles.css), accessibility stylesheet (a11y.css) and '
    'site JavaScript (site.js) were also reviewed and updated.',
    space_after=8)

add_rule(doc, '0A9E9E', 6)

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 2 — ISSUES RESOLVED
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '2.  Issues resolved', level=1)

# ── 2a Design & Visual ────────────────────────────────────────────────────────
add_heading(doc, 'Design & Visual', level=2, space_before=10)

add_tick(doc, '"SERVICE" label removed — the word "Service" was appearing as an eyebrow label above the page title in the hero section on the landing page and the landlord page. It added no information and has been removed.')
add_tick(doc, 'Crosshair/target SVG removed — a decorative SVG in the hero section rendered as a target or crosshair symbol, which was inappropriate for a housing licensing service. It has been removed from both hero sections.')
add_tick(doc, '"in Gloucester" text wrapping fixed — on most pages the words "in Gloucester" were orphaning to a new line in the hero description. A no-wrap fix has been applied so the phrase always stays on one line.')
add_tick(doc, 'Tab scrollbar removed — the tab navigation bar was showing a thin scrollbar on desktop. The tab container now hides the scrollbar while still allowing horizontal scroll on mobile touch devices.')
add_tick(doc, 'Regulation list numbering fixed — the management duties section on the Standards page used an auto-numbered list that conflicted with the regulation numbers (Regulation 3, Regulation 4, etc.), creating confusing labels like "1. Regulation 3". The automatic numbering has been removed; only the regulation numbers remain.')

# ── 2b Naming ─────────────────────────────────────────────────────────────────
add_heading(doc, 'Naming consistency', level=2, space_before=10)
add_body(doc,
    'An agreed naming convention has been applied across all pages:',
    space_after=3)
add_body(doc, '·  First mention per page: Houses in Multiple Occupation (HMO)', space_after=2)
add_body(doc, '·  All subsequent mentions: HMO', space_after=2)
add_body(doc, '·  Formal legislation titles (e.g. "The Management of Houses in Multiple Occupation (England) Regulations 2006") are always written in full.', space_after=6)

add_tick(doc, 'Planning page title, h1 and breadcrumb updated to "Planning permission for Houses in Multiple Occupation (HMO)".')
add_tick(doc, 'Body text on the planning page corrected — "House in Multiple Occupation (HMO)" replaced with "HMO" after the first use, and "Houses in Multiple Occupation" corrected to "HMOs" in the corporate plan commitment text.')
add_tick(doc, 'Landlord page lede updated — "Houses in Multiple Occupation" replaced with "HMO" as it appeared after the first mention.')
add_tick(doc, 'Grammar fix — "a HMO" corrected to "an HMO" on the licensing page.')

# ── 2c Navigation & Structure ─────────────────────────────────────────────────
add_heading(doc, 'Navigation & Structure', level=2, space_before=10)

add_tick(doc, 'AI-generated content removed — a "Looking for something else?" section on the landing page and landlord page could not be traced to the master data schema and appeared to have been auto-generated. It also duplicated an existing planning permissions link. Both sections have been removed.')
add_tick(doc, 'Duplicate "On this page" sidebar cards removed — on pages that already have a tab bar at the top (Licensing, Residents, Standards, Planning), the sidebar contained an "On this page" box with identical links. This duplication has been removed; the tab bar is the single navigation mechanism.')
add_tick(doc, 'FAQ content distributed — the standalone FAQ page contained questions that belonged in context with specific sections of the site. Resident and tenant questions have been embedded at the bottom of the Residents page; landlord and enforcement questions have been embedded at the bottom of the Landlord page.')
add_tick(doc, 'FAQ page converted — the FAQ page now serves as a short navigation page directing users to the embedded question sections, rather than duplicating content.')
add_tick(doc, 'Service card link titles corrected — card headings on the landing page now match the h1 of the page they link to. "I\'m a resident in an HMO" now reads "Living in an HMO — your rights and what to expect". "I\'m a landlord or manager" now reads "HMO landlord and manager information".')
add_tick(doc, 'Page lede duplication fixed — on the Apply page and Residents page, the subtitle was restating the page title in different words. Both ledes have been rewritten to add new, useful information.')

# ── 2d Functionality & Links ──────────────────────────────────────────────────
add_heading(doc, 'Functionality & Links', level=2, space_before=10)

add_tick(doc, '"Check if an HMO is licensed" link fixed — was pointing to "#" (broken). Now links to the public HMO licence register on gloucester.gov.uk.')
add_tick(doc, '"Report a problem in an HMO" link fixed — was pointing to "#" (broken). Now links to the council\'s online contact form.')
add_tick(doc, 'Footer links updated — all page footers previously linked to the old standalone FAQ page. These now link to the Landlord information page or other relevant destinations.')

# ── 2e UI Components ──────────────────────────────────────────────────────────
add_heading(doc, 'UI Components', level=2, space_before=10)

add_tick(doc, 'Accordion behaviour fixed — when opening a question in an accordion/FAQ section, any previously open question now closes automatically. Previously, multiple questions could be open at once.')
add_tick(doc, 'Back to top added — a "Back to top" link has been added at the bottom of every page, making it easier to navigate back up on longer pages.')

# ── 2f Accessibility ──────────────────────────────────────────────────────────
add_heading(doc, 'Accessibility', level=2, space_before=10)

add_tick(doc, 'Accessibility toolbar icon replaced — the previous icon (a stick figure with arms out) was not clearly recognisable as an accessibility tool. It has been replaced with an eye/view icon that more clearly communicates the purpose of the button.')
add_tick(doc, 'Reading ruler performance improved — when using the strip ruler accessibility tool, there was noticeable lag as the ruler tracked the mouse. The update logic has been wrapped in requestAnimationFrame, which reduces lag by syncing updates to the browser\'s render cycle.')
add_tick(doc, 'Dark theme search box contrast fixed — in dark mode, the search input field was almost invisible because its background matched the page background. The input now uses a white background in dark mode.')
add_tick(doc, 'Dark + High Contrast mode fixes — several elements were not displaying correctly when both dark theme and high contrast were enabled simultaneously: the Popular Tasks section heading was invisible, table text was unreadable, and button text disappeared on hover. All have been corrected with targeted CSS overrides.')
add_tick(doc, 'Button hover state fixed — in certain theme combinations, hovering over a primary button caused the button text to disappear. The colour rules have been clarified so text is always visible regardless of theme.')

add_rule(doc)

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 3 — DECISIONS MADE
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '3.  Decisions made', level=1)
add_body(doc, 'The following decisions were made during the session and are recorded here for the team\'s reference.', space_after=6)

decisions = [
    ('HMO naming convention',
     'Full form "Houses in Multiple Occupation (HMO)" on first mention per page. "HMO" for all subsequent uses. Formal legislation titles always in full.'),
    ('Landing page',
     'Keep the landing page but simplify it. The "Looking for something else?" AI-generated section and the "Service" label were removed. The service card grid and popular tasks remain.'),
    ('FAQ page',
     'Distribute FAQ content into the relevant pages (residents / landlords). Convert faq.html to a short navigation page rather than deleting it, to avoid breaking any bookmarks or existing links.'),
    ('"On this page" vs Tabs',
     'On pages that have a tab bar, the "On this page" sidebar card is a duplicate and has been removed. The tab bar is the single in-page navigation.'),
    ('"Check if an HMO is licensed" link',
     'Links to the public HMO licence register spreadsheet on gloucester.gov.uk.'),
    ('"Report a problem in an HMO" link',
     'Links to the AchieveService contact form at gloucester-self.achieveservice.com.'),
    ('"Renew an HMO licence" popular task',
     'Renewal is handled through the same application form as a new licence. The link correctly points to apply.html.'),
]

tbl = doc.add_table(rows=len(decisions)+1, cols=2)
tbl.style = 'Table Grid'
# Header
for i, txt in enumerate(['Decision', 'Outcome']):
    cell = tbl.cell(0, i)
    shade_cell(cell, '0054A4')
    run = cell.paragraphs[0].add_run(txt)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)
for r, (d, o) in enumerate(decisions, 1):
    c0 = tbl.cell(r, 0)
    c1 = tbl.cell(r, 1)
    rn = c0.paragraphs[0].add_run(d)
    rn.font.size = Pt(10)
    rn.bold = True
    c1.paragraphs[0].add_run(o).font.size = Pt(10)
    if r % 2 == 0:
        shade_cell(c0, 'F1F3F5')
        shade_cell(c1, 'F1F3F5')

doc.add_paragraph()
add_rule(doc, '0A9E9E', 6)

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 4 — OUTSTANDING WORK
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '4.  Outstanding work', level=1)
add_body(doc,
    'The following items have been logged as TODOs. Priority is indicated for each. '
    'High priority items should be completed before the branch is merged.',
    space_after=8)

# Priority key
tbl_key = doc.add_table(rows=1, cols=3)
for i, (label, col_hex, desc) in enumerate([
    ('High', '0054A4', 'Must be done before merge'),
    ('Medium', '0A9E9E', 'Should be done before launch'),
    ('Low', '6C757D', 'Can be addressed post-launch'),
]):
    cell = tbl_key.cell(0, i)
    shade_cell(cell, col_hex)
    p = cell.paragraphs[0]
    r1 = p.add_run(f'{label}  ')
    r1.bold = True
    r1.font.color.rgb = WHITE
    r1.font.size = Pt(9.5)
    r2 = p.add_run(desc)
    r2.font.color.rgb = WHITE
    r2.font.size = Pt(9.5)
doc.add_paragraph()

todos = [
    ('High',   'TODO-01  Mobile layout review',
               'Check all HMO pages at 375px and 768px viewport widths. Known concern: the layout '
               'is considered cluttered on small screens. Focus areas: service card grid, tab bar overflow, '
               'hero padding, popular tasks grid, sidebar stacking order, and back-to-top positioning.'),
    ('High',   'TODO-02  Full accessibility audit',
               'Run a structured audit across all HMO pages covering WCAG 2.2 AA colour contrast '
               '(all theme and contrast combinations), keyboard navigation, screen reader semantics, '
               'touch target sizes, and form accessibility on the Apply page. Report and fix any failures.'),
    ('Medium', 'TODO-03  Footer colour token alignment',
               'The footer uses a darker blue (#002650) than the header and hero (#0054A4). The team '
               'needs to decide whether to align them to a consistent tone or confirm the contrast is '
               'intentional. Whichever colour is chosen, text and link contrast must be verified for '
               'all themes.'),
    ('Medium', 'TODO-07  Spacing inconsistencies sweep',
               'A full visual spacing review was deferred. Items to check: padding above and below the '
               'hero, gap between the tab bar and first section heading, sidebar card spacing on narrow '
               'viewports, and bottom padding on the Apply page.'),
    ('Low',    'TODO-04  Tone and naming sweep (non-HMO templates)',
               'The HMO pages have been updated but the shared templates and other service examples '
               '(Beta-website-sample-v1, Taxi website, etc.) have not been reviewed. These should be '
               'checked for the same naming, capitalisation and terminology issues.'),
    ('Low',    'TODO-05  Overall design modernisation',
               'The current design meets government standards but the original brief called for a more '
               'modern, contemporary aesthetic. Proposed improvements discussed but not yet implemented '
               'include: hero gradient or clip-path, tighter letter-spacing on display headings, and '
               'animated hover states on popular tasks. This requires a separate design conversation '
               'before implementation.'),
    ('Low',    'TODO-06  "in Gloucester" wrapping on remaining pages',
               'The no-wrap fix was applied to the landing page and landlord page hero ledes. The same '
               'issue may occur on other pages depending on viewport width. A quick sweep should confirm '
               'whether any other ledescontain the same orphan.'),
]

for priority, label, description in todos:
    col = GCC_BLUE if priority == 'High' else (TEAL if priority == 'Medium' else MID_GREY)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(1)
    r_pri = p.add_run(f'[{priority}]  ')
    r_pri.bold = True
    r_pri.font.color.rgb = col
    r_pri.font.size = Pt(10.5)
    r_lbl = p.add_run(label)
    r_lbl.bold = True
    r_lbl.font.size = Pt(10.5)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.7)
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(description)
    r2.font.size = Pt(10)
    r2.font.color.rgb = MID_GREY

doc.add_paragraph()
add_rule(doc)

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 5 — TECHNICAL REFERENCE
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, '5.  Technical reference', level=1)
add_body(doc, 'For developers picking up this work:', space_after=4)

ref = [
    ('Repository',      'Gloucester-City-Council/GCC-Beta-site-v1'),
    ('Branch',          'claude/elegant-knuth-v8YaS'),
    ('Pull request',    'PR #43 — https://github.com/Gloucester-City-Council/GCC-Beta-site-v1/pull/43'),
    ('Context file',    'CLAUDE.md at the repository root — full session log, decisions, completed work and TODOs'),
    ('HMO pages',       '/HMO/web/*.html'),
    ('Stylesheet',      '/templates/assets/styles.css'),
    ('JavaScript',      '/templates/assets/site.js'),
    ('Master schema',   '/HMO/hmo_master_schema_v1.json'),
    ('Files changed',   '10 files — 362 insertions, 337 deletions'),
]
tbl = doc.add_table(rows=len(ref), cols=2)
tbl.style = 'Table Grid'
for r, (k, v) in enumerate(ref):
    c0 = tbl.cell(r, 0)
    c1 = tbl.cell(r, 1)
    rn0 = c0.paragraphs[0].add_run(k)
    rn0.bold = True
    rn0.font.size = Pt(10)
    c1.paragraphs[0].add_run(v).font.size = Pt(10)
    if r % 2 == 0:
        shade_cell(c0, 'F1F3F5')
        shade_cell(c1, 'F1F3F5')

doc.add_paragraph()
add_body(doc,
    'Note: this is a purely static site (HTML, CSS and vanilla JavaScript). There is no build step, '
    'no framework, and no package manager. Files can be edited directly and previewed in any browser.',
    space_after=8, italic=True, colour=MID_GREY)

add_rule(doc, '0054A4', 12)

# ── Footer ────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'Gloucester City Council Digital Team  ·  {datetime.date.today().strftime("%-d %B %Y")}')
r.font.size = Pt(9)
r.font.color.rgb = MID_GREY

# ── Save ──────────────────────────────────────────────────────────────────────
out = '/home/user/GCC-Beta-site-v1/GCC_HMO_Beta_Development_Summary.docx'
doc.save(out)
print(f'Saved: {out}')
